import os
import re
import pandas as pd
from docx import Document
import google.generativeai as genai
from google.generativeai import GenerativeModel
import undetected_chromedriver as uc
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
import csv
import requests
import json
import random
import shutil
from PIL import Image, ImageEnhance
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
import configparser
from pathlib import Path

class ConfigManager:
    def __init__(self):
        self.config_file = Path.home() / ".article_automation_config.ini"
        self.config = configparser.ConfigParser()
        self.load_config()
    
    def load_config(self):
        """Load configuration from file"""
        if self.config_file.exists():
            self.config.read(self.config_file)
        else:
            self.set_default_config()
    
    def set_default_config(self):
        """Set default configuration values"""
        self.config['credentials'] = {
            'username': '',
            'password': '',
            'url': 'https://pharmastan.net/pharma-login'
        }
        
        self.config['paths'] = {
            'excel': '',
            'word': '',
            'failed': '',
            'done': '',
            'logo_image_path': '',
            'original_images_path': '',
            'featured_images_path': ''
        }
        
        self.config['timeouts'] = {
            'page_load': '60000',
            'element': '10000',
            'hover': '1000',
            'click': '2000',
            'typing': '120'
        }
        
        self.config['apis'] = {
            'gemini': '',
            'fireworks': ''
        }
    
    def save_config(self):
        """Save configuration to file"""
        with open(self.config_file, 'w') as f:
            self.config.write(f)
    
    def get_config_dict(self):
        """Convert config to dictionary format used by the main program"""
        return {
            "credentials": {
                "username": self.config.get('credentials', 'username'),
                "password": self.config.get('credentials', 'password'),
                "url": self.config.get('credentials', 'url')
            },
            "paths": {
                "excel": self.config.get('paths', 'excel'),
                "word": self.config.get('paths', 'word'),
                "failed": self.config.get('paths', 'failed'),
                "done": self.config.get('paths', 'done'),
                "logo_image_path": self.config.get('paths', 'logo_image_path'),
                "original_images_path": self.config.get('paths', 'original_images_path'),
                "featured_images_path": self.config.get('paths', 'featured_images_path')
            },
            "timeouts": {
                "page_load": int(self.config.get('timeouts', 'page_load')),
                "element": int(self.config.get('timeouts', 'element')),
                "hover": int(self.config.get('timeouts', 'hover')),
                "click": int(self.config.get('timeouts', 'click')),
                "typing": int(self.config.get('timeouts', 'typing'))
            },
            "APIs": {
                "gemini": self.config.get('apis', 'gemini'),
                "firworks": self.config.get('apis', 'fireworks')
            }
        }

class DocxToExcelConverter:
    def __init__(self):
        self.paragraphs = []
        self.articles_data = {}
        self.excel_paths = []

    def read_docx(self, docx_path):
        doc = Document(docx_path)
        self.paragraphs = doc.paragraphs

    def get_actual_bullet_or_number(self, paragraph):
        element = paragraph._element
        numPr = element.find('.//w:numPr', namespaces=element.nsmap)
        if numPr is None:
            return None

        bullet_char = '•'
        lvl = element.find('.//w:ilvl', namespaces=element.nsmap)
        if lvl is not None:
            return bullet_char + ' '

        numId = element.find('.//w:numId', namespaces=element.nsmap)
        if numId is not None:
            return "1. "

        return None

    def get_clean_text_with_formatting(self):
        lines = []

        for para in self.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue

            prefix = self.get_actual_bullet_or_number(para)
            if prefix:
                lines.append(prefix + text)
            else:
                pPr = para._element.find('.//w:pPr', namespaces=para._element.nsmap)
                if pPr is not None:
                    ind = pPr.find('.//w:ind', namespaces=para._element.nsmap)
                    if ind is not None:
                        left = ind.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left')
                        if left:
                            lines.append('    ' * (int(left) // 720) + text)
                            continue
                lines.append(text)
        return "\n".join(lines)

    def sanitize_filename(self, name):
        return re.sub(r'[\\/*?:"<>|]', "_", name)

    def extract_structure(self):
        self.articles_data = {}
        full_text = self.get_clean_text_with_formatting()

        article_pattern = r'###\s*(.*?)\s*###'
        article_matches = list(re.finditer(article_pattern, full_text, re.DOTALL))

        for i, article_match in enumerate(article_matches):
            article_title = article_match.group(1).strip()
            start = article_match.end()
            end = article_matches[i + 1].start() if i + 1 < len(article_matches) else len(full_text)
            article_content = full_text[start:end].strip()

            subsection_pattern = r'\${3}\s*(.*?)\s*\${3}'
            subsection_matches = list(re.finditer(subsection_pattern, article_content, re.DOTALL))

            sections = []
            for j, sub_match in enumerate(subsection_matches):
                subsection_title = sub_match.group(1).strip()
                body_start = sub_match.end()
                body_end = subsection_matches[j + 1].start() if j + 1 < len(subsection_matches) else len(article_content)
                subsection_body = article_content[body_start:body_end].strip()
                subsection_body = re.sub(r'\n{2,}', '\n', subsection_body)
                sections.append([article_title, subsection_title, subsection_body])

            self.articles_data[article_title] = sections

    def save_each_article_to_excel(self, output_folder):
        os.makedirs(output_folder, exist_ok=True)
        for article_title, sections in self.articles_data.items():
            safe_title = self.sanitize_filename(article_title)
            output_path = os.path.join(output_folder, f"{safe_title}.xlsx")
            df = pd.DataFrame(sections, columns=["Article Title", "Subsection Title", "Content"])
            self.excel_paths.append(output_path)            
            writer = pd.ExcelWriter(output_path, engine='xlsxwriter')
            df.to_excel(writer, index=False, sheet_name='Content')
            workbook = writer.book
            worksheet = writer.sheets['Content']
            worksheet.set_column('A:A', 30)
            worksheet.set_column('B:B', 30)
            worksheet.set_column('C:C', 80)
            text_format = workbook.add_format({'text_wrap': True})
            worksheet.set_column('C:C', None, text_format)
            writer.close()

    def convert(self, input_folder, output_folder):
        if not os.path.exists(input_folder):
            print(f"❌ Input folder not found: {input_folder}")
            return
        os.makedirs(output_folder, exist_ok=True)
        for filename in os.listdir(input_folder):
            if filename.lower().endswith('.docx'):
                full_path = os.path.join(input_folder, filename)
                print(f"📄 Processing: {filename}")
                self.read_docx(full_path)
                self.extract_structure()
                base_name = os.path.splitext(filename)[0]
                safe_folder = self.sanitize_filename(base_name)
                article_output_folder = os.path.join(output_folder, safe_folder)
                print(f"📂 Saving to: {article_output_folder}")
                self.save_each_article_to_excel(article_output_folder)
                shutil.move(full_path, CONFIG['paths']['done'])
        print("✅ All Word files processed and Excel outputs created.")

def convert_to_paragraphs(text: str) -> str:
    """Convert Arabic text with line breaks into HTML paragraphs."""
    lines = text.splitlines()
    paragraphs = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r"^[\-\•\*]\s+", stripped):
            bullet_content = re.sub(r"^[\-\•\*]\s*", "", stripped)
            paragraphs.append(
                '<p dir="rtl" style="text-indent:-1.5em; padding-right:2em;">'
                '&bull;&emsp;' + bullet_content +
                '</p>'
            )
        else:
            paragraphs.append(f'<p dir="rtl">{stripped}</p>')

    return ''.join(paragraphs)

def generate_arabic_meta_expert_summary_fireworks(excel_path, api_key):
    """Generate a 2-line Arabic meta description from an Excel file using Fireworks AI."""
    url = "https://api.fireworks.ai/inference/v1/completions"
 
    df = pd.read_excel(excel_path)
    df.fillna("", inplace=True)

    article_title = df.iloc[0, 0]
    article_intro = df.iloc[0, 2]

    subsections = []
    for _, row in df.iterrows():
        subtitle = row.iloc[1]
        body = row.iloc[2]
        if subtitle or body:
            subsections.append(f"{subtitle}\n{body}".strip())
    full_content = f"{article_intro.strip()}\n\n" + "\n\n".join(subsections)

    prompt = (
        f"أنشئ وصفًا ميتا احترافيًا باللغة العربية لا يتجاوز 160 حرفًا،"
        f" دون شرح أو خطوات، ويكون موجزًا ومناسبًا لمحركات البحث (SEO)،"
        f" ويحتوي على كلمات مفتاحية طبية مناسبة لموضوع المقال التالي:\n\n"
        f"العنوان: {article_title}\n\n"
        f"المحتوى:\n{full_content}\n\n"
        f"الوصف الميتا:"
    )
    payload = {
        "model": "accounts/fireworks/models/llama4-maverick-instruct-basic",
        "max_tokens": 70,
        "top_p": 1,
        "top_k": 40,
        "temperature": 0.1,
        "presence_penalty": 0,
        "frequency_penalty": 0,
        "stop": ["\n", "انتهى", "الخطوة", "هل يمكن"],
        "prompt": prompt
    }
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    response = requests.post(url, headers=headers, data=json.dumps(payload))
    if response.status_code == 200:
        result = response.json()
        meta_text = result.get("choices", [{}])[0].get("text", "").strip()
        if not meta_text:
            raise ValueError("Fireworks response is empty")
        return {
            'meta_expert_summary': meta_text
        }
    else:
        print(f"❌ API Error: {response.status_code} - {response.text}")
        raise Exception(f"API request failed with status code {response.status_code}")

def generate_arabic_meta_expert_summary(excel_path, gemini_api_key):
    genai.configure(api_key=gemini_api_key)
    model = GenerativeModel('gemini-2.0-flash-lite-001')
    df = pd.read_excel(excel_path)
    df.fillna("", inplace=True)
    
    article_title = df.iloc[0,0]
    article_intro = df.iloc[0,2]
    
    subsections = []
    for _, row in df.iterrows():
        title = row.iloc[1] 
        body = row.iloc[2] 
        if title or body:
            subsections.append(f"{title}\n{body}")

    full_content = f"{article_intro}\n\n" + "\n\n".join(subsections)

    prompt = f"""
المقال التالي يتحدث عن موضوع طبي بعنوان "{article_title}".

المحتوى:
{full_content}

🧠 المطلوب:
اكتب ملخصًا احترافيًا موجزًا (سطرين فقط كحد أقصى) يصلح كوصف ميتا للمقال. يجب أن يكون باللغة العربية ويعبر بدقة عن محتوى المقال بطريقة جذابة وواضحة.
    """
    response = model.generate_content(prompt)
    if response.text.strip() == "":
        raise ValueError("gemini response is empty")

    return {
        'meta_expert_summary': response.text.strip()
    }

def add_watermark_from_folder(folder_path, output_path, watermark_image_path, base_size=(616, 367), transparency=0.3, allowed_extensions=(".jpg", ".jpeg", ".png")):
    image_files = [
        f for f in os.listdir(folder_path)
        if f.lower().endswith(allowed_extensions)
    ]
    if not image_files:
        raise FileNotFoundError("No image files found in the folder.")
    
    chosen_image = random.choice(image_files)
    base_image_path = os.path.join(folder_path, chosen_image)
    
    base_img = Image.open(base_image_path).convert("RGBA")
    base_img = base_img.resize(base_size)
    
    watermark = Image.open(watermark_image_path).convert("RGBA")
    watermark = watermark.resize(base_size)
    
    alpha = watermark.split()[3]
    alpha = ImageEnhance.Brightness(alpha).enhance(transparency)
    watermark.putalpha(alpha)

    combined = Image.alpha_composite(base_img, watermark)
    
    combined.convert("RGB").save(os.path.join(output_path, chosen_image), "JPEG")
    return os.path.join(output_path, chosen_image)

def login_to_website(config):
    options = uc.ChromeOptions()
    prefs = {
        "credentials_enable_service": False,
        "profile.password_manager_enabled": False
    }
    options.add_experimental_option("prefs", prefs)
    
    driver = uc.Chrome(options=options)
    wait = WebDriverWait(driver, 15)
    
    driver.get(config["credentials"]["url"])
    
    wait.until(EC.presence_of_element_located((By.ID, "user_login"))).send_keys(config["credentials"]["username"])
    wait.until(EC.presence_of_element_located((By.ID, "user_pass"))).send_keys(config["credentials"]["password"])
    wait.until(EC.element_to_be_clickable((By.ID, "wp-submit"))).click()
    wait.until(EC.url_contains("wp-admin"))

    return driver

def fill_article_intro(driver, text):
    text = convert_to_paragraphs(text)
    wait = WebDriverWait(driver, 15)
    textarea = wait.until(EC.visibility_of_element_located((By.ID, "content")))
    textarea.clear()
    textarea.send_keys(text)
    driver.switch_to.default_content()
    return driver

def click_add_another_and_wait(driver):
    wait = WebDriverWait(driver, 15)
    initial_count = len(driver.find_elements(By.NAME, "post_contents_desc[]"))
    add_button = wait.until(EC.element_to_be_clickable((By.ID, "add-contents-row")))
    add_button.click()
    wait.until(lambda d: len(d.find_elements(By.NAME, "post_contents_desc[]")) > initial_count)

def click_all_html_toggle_buttons(driver):
    wait = WebDriverWait(driver, 15)

    html_buttons = driver.find_elements(By.TAG_NAME, "button")
    for i, btn in enumerate(html_buttons):
        btn_id = btn.get_attribute("id")
        if btn_id and btn_id.endswith("-html") and "post_contents_desc" in btn_id:
            try:
                button = wait.until(EC.presence_of_element_located((By.ID, btn_id)))
                driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", button)

                if button.get_attribute("aria-pressed") != "true":
                    wait.until(EC.element_to_be_clickable((By.ID, btn_id))).click()
                    time.sleep(0.2)
            except Exception as e:
                print(f"[{i}] Failed to click button {btn_id}: {e}")

def fill_all_subsections_bodies(driver, content_list):
    wait = WebDriverWait(driver, 15)
    count = len(driver.find_elements(By.NAME, "post_contents_desc[]"))
    print(f"{count} textareas found. Filling them...")
    for i in range(count):
        textarea = driver.find_element(By.ID, f"post_contents_desc{i}")
        wait.until(EC.visibility_of(textarea))
        content = content_list[i] if i < len(content_list) else f"محتوى افتراضي للجزء {i+1}"
        textarea.clear()
        textarea.send_keys(convert_to_paragraphs(content))

def click_save_draft(driver):
    wait = WebDriverWait(driver, 15)
    save_button = wait.until(EC.element_to_be_clickable((By.ID, "save-post")))
    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", save_button)
    save_button.click()

def click_publish(driver):
    wait = WebDriverWait(driver, 20)

    try:
        wait.until(lambda d: d.execute_script("return document.readyState") == "complete")

        publish_button = wait.until(EC.element_to_be_clickable((By.ID, "publish")))

        driver.execute_script("arguments[0].scrollIntoView(true);", publish_button)
        time.sleep(0.5)

        driver.execute_script("arguments[0].click();", publish_button)

        wait.until(EC.visibility_of_element_located(
            (By.XPATH, "//p[contains(text(), 'Post published.')]")
        ))

        print("✅ Article published successfully.")

    except Exception as e:
        print("❌ Failed to publish:", e)

def get_category_id_by_name(category_name, csv_path="disease_categories.csv"):
    with open(csv_path, newline='', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        for row in reader:
            if row["category_name"].strip() == category_name.strip():
                return row["id"]
    raise ValueError(f"Category '{category_name}' not found in CSV.")

def select_category_by_name(driver, category_name, csv_path="disease_categories.csv"):
    category_id = get_category_id_by_name(category_name, csv_path)
    wait = WebDriverWait(driver, 10)
    
    selector = f"input[name='tax_input[diseases_category][]'][value='{category_id}']"
    checkbox = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, selector)))

    if not checkbox.is_selected():
        driver.execute_script("arguments[0].click();", checkbox)

def upload_featured_image(driver, image_path, timeout=30):
    wait = WebDriverWait(driver, timeout)
    set_thumb_button = wait.until(EC.element_to_be_clickable((By.ID, "set-post-thumbnail")))
    driver.execute_script("arguments[0].click();", set_thumb_button)

    wait = WebDriverWait(driver, timeout)
    upload_tab = wait.until(EC.element_to_be_clickable((By.ID, "menu-item-upload")))
    upload_tab.click()

    file_input = driver.find_element(By.CSS_SELECTOR, "input[type='file']")
    file_input.send_keys(image_path)
    time.sleep(2)

    button = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable(
            (By.CSS_SELECTOR, "button.media-button-select")
        )
    )
    button.click()
    print("Button clicked successfully!")

def press_all_code_editors(driver):
    wait = WebDriverWait(driver, 15)
    buttons = wait.until(lambda d: d.find_elements(By.CSS_SELECTOR, "button.wp-switch-editor.switch-html"))
    for i, button in enumerate(buttons):
        try:
            current_button = wait.until(lambda d: d.find_elements(By.CSS_SELECTOR, "button.wp-switch-editor.switch-html"))[i]
            wait.until(lambda d: current_button.is_displayed() and current_button.is_enabled())
            
            if current_button.get_attribute("aria-pressed") != "true":
                driver.execute_script("""
                    arguments[0].scrollIntoView({
                        behavior: 'auto',
                        block: 'center',
                        inline: 'center'
                    });
                    arguments[0].click();
                """, current_button)
                print(f"✅ Switched editor #{i} to Code mode.")
            else:
                print(f"⏩ Editor #{i} already in Code mode.")
        except Exception as e:
            print(f"❌ Failed to switch editor #{i}: {e}")

def article_writter(excel_path, config):
    article = pd.read_excel(excel_path)
    category = os.path.basename(os.path.dirname(excel_path))
    title = article.iloc[0,0]
    intro = article.iloc[0,2]
    subtitles = article.iloc[1:,1].tolist()
    subtitles_bodies = article.iloc[1:,2].tolist()
    num_subsections = article.shape[0]-1

    driver = login_to_website(config)
    driver.get("https://pharmastan.net/wp-admin/post-new.php?post_type=disease")
    wait = WebDriverWait(driver, 15)
    wait.until(EC.visibility_of_element_located((By.ID, "title"))).send_keys(title)

    wait = WebDriverWait(driver, 10)
    code_btn = wait.until(EC.element_to_be_clickable((By.ID, "content-html")))
    code_btn.click()
    time.sleep(2)
    driver = fill_article_intro(driver,intro)
    time.sleep(2)

    button = WebDriverWait(driver, 10).until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, "button.switch-tmce"))
    )
    if button.get_attribute("aria-pressed") != "true":
        button.click()
    
    press_all_code_editors(driver)
    
    for i in range(num_subsections-1):
        click_add_another_and_wait(driver)
    subsections_titles = driver.find_elements(By.NAME, "post_content_titles[]")
    
    for i, field in enumerate(subsections_titles):
            field.clear()
            field.send_keys(subtitles[i])
    press_all_code_editors(driver)
    
    fill_all_subsections_bodies(driver,subtitles_bodies)
    select_category_by_name(driver, category)
    textarea = driver.find_element(By.ID, "excerpt")
    textarea.clear()
    try:
        excerpt_text = generate_arabic_meta_expert_summary(excel_path, config['APIs']['gemini'])['meta_expert_summary']
    except:
        excerpt_text = generate_arabic_meta_expert_summary_fireworks(excel_path, config['APIs']['firworks'])['meta_expert_summary']
    textarea.send_keys(excerpt_text)
    time.sleep(2)
    featured_img_path = add_watermark_from_folder(
        folder_path=config['paths']['original_images_path'],
        output_path=config['paths']['featured_images_path'],
        watermark_image_path=config['paths']['logo_image_path'],
        base_size=(616, 367),
        transparency=0.2
    )
    upload_featured_image(driver,featured_img_path)
    time.sleep(2)
    click_publish(driver)

class ArticleAutomationGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Article Automation Tool")
        self.root.geometry("800x700")
        
        self.config_manager = ConfigManager()
        self.create_widgets()
        self.load_saved_config()
        
    def create_widgets(self):
        # Create notebook for tabs
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Configuration tab
        config_frame = ttk.Frame(notebook)
        notebook.add(config_frame, text="Configuration")
        
        # Credentials section
        cred_frame = ttk.LabelFrame(config_frame, text="Website Credentials", padding=10)
        cred_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Label(cred_frame, text="Username:").grid(row=0, column=0, sticky='w', padx=5, pady=2)
        self.username_var = tk.StringVar()
        ttk.Entry(cred_frame, textvariable=self.username_var, width=50).grid(row=0, column=1, padx=5, pady=2)
        
        ttk.Label(cred_frame, text="Password:").grid(row=1, column=0, sticky='w', padx=5, pady=2)
        self.password_var = tk.StringVar()
        ttk.Entry(cred_frame, textvariable=self.password_var, show="*", width=50).grid(row=1, column=1, padx=5, pady=2)
        
        ttk.Label(cred_frame, text="Login URL:").grid(row=2, column=0, sticky='w', padx=5, pady=2)
        self.url_var = tk.StringVar()
        ttk.Entry(cred_frame, textvariable=self.url_var, width=50).grid(row=2, column=1, padx=5, pady=2)
        
        # Paths section
        paths_frame = ttk.LabelFrame(config_frame, text="File Paths", padding=10)
        paths_frame.pack(fill='x', padx=5, pady=5)
        
        self.path_vars = {}
        path_labels = [
            ("Excel Output Path:", "excel"),
            ("Word Input Path:", "word"),
            ("Failed Files Path:", "failed"),
            ("Done Files Path:", "done"),
            ("Logo Image Path:", "logo_image_path"),
            ("Original Images Path:", "original_images_path"),
            ("Featured Images Path:", "featured_images_path")
        ]
        
        for i, (label, key) in enumerate(path_labels):
            ttk.Label(paths_frame, text=label).grid(row=i, column=0, sticky='w', padx=5, pady=2)
            self.path_vars[key] = tk.StringVar()
            ttk.Entry(paths_frame, textvariable=self.path_vars[key], width=40).grid(row=i, column=1, padx=5, pady=2)
            if key == "logo_image_path":
                ttk.Button(paths_frame, text="Browse", 
                          command=lambda k=key: self.browse_file(k)).grid(row=i, column=2, padx=5, pady=2)
            else:
                ttk.Button(paths_frame, text="Browse", 
                          command=lambda k=key: self.browse_folder(k)).grid(row=i, column=2, padx=5, pady=2)
        
        # API Keys section
        api_frame = ttk.LabelFrame(config_frame, text="API Keys", padding=10)
        api_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Label(api_frame, text="Gemini API Key:").grid(row=0, column=0, sticky='w', padx=5, pady=2)
        self.gemini_var = tk.StringVar()
        ttk.Entry(api_frame, textvariable=self.gemini_var, width=50, show="*").grid(row=0, column=1, padx=5, pady=2)
        
        ttk.Label(api_frame, text="Fireworks API Key:").grid(row=1, column=0, sticky='w', padx=5, pady=2)
        self.fireworks_var = tk.StringVar()
        ttk.Entry(api_frame, textvariable=self.fireworks_var, width=50, show="*").grid(row=1, column=1, padx=5, pady=2)
        
        # Timeouts section
        timeout_frame = ttk.LabelFrame(config_frame, text="Timeouts (ms)", padding=10)
        timeout_frame.pack(fill='x', padx=5, pady=5)
        
        self.timeout_vars = {}
        timeout_labels = [
            ("Page Load:", "page_load"),
            ("Element:", "element"),
            ("Hover:", "hover"),
            ("Click:", "click"),
            ("Typing:", "typing")
        ]
        
        for i, (label, key) in enumerate(timeout_labels):
            ttk.Label(timeout_frame, text=label).grid(row=i//3, column=(i%3)*2, sticky='w', padx=5, pady=2)
            self.timeout_vars[key] = tk.StringVar()
            ttk.Entry(timeout_frame, textvariable=self.timeout_vars[key], width=10).grid(row=i//3, column=(i%3)*2+1, padx=5, pady=2)
        
        # Save Configuration button
        ttk.Button(config_frame, text="Save Configuration", 
                  command=self.save_configuration).pack(pady=10)
        
        # Processing tab
        process_frame = ttk.Frame(notebook)
        notebook.add(process_frame, text="Processing")
        
        # Processing controls
        control_frame = ttk.LabelFrame(process_frame, text="Processing Controls", padding=10)
        control_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Button(control_frame, text="Start Processing", 
                  command=self.start_processing).pack(side='left', padx=5)
        ttk.Button(control_frame, text="Stop Processing", 
                  command=self.stop_processing).pack(side='left', padx=5)
        
        # Progress bar
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(process_frame, variable=self.progress_var, 
                                          maximum=100, length=400)
        self.progress_bar.pack(pady=10)
        
        # Status label
        self.status_var = tk.StringVar(value="Ready")
        ttk.Label(process_frame, textvariable=self.status_var).pack(pady=5)
        
        # Log text area
        log_frame = ttk.LabelFrame(process_frame, text="Processing Log", padding=10)
        log_frame.pack(fill='both', expand=True, padx=5, pady=5)
        
        self.log_text = tk.Text(log_frame, height=15, width=80)
        scrollbar = ttk.Scrollbar(log_frame, orient='vertical', command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=scrollbar.set)
        
        self.log_text.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
        self.processing_thread = None
        self.stop_flag = threading.Event()
    
    def browse_folder(self, key):
        folder = filedialog.askdirectory()
        if folder:
            self.path_vars[key].set(folder)
    
    def browse_file(self, key):
        file_path = filedialog.askopenfilename(
            filetypes=[("Image files", "*.png *.jpg *.jpeg *.gif *.bmp")]
        )
        if file_path:
            self.path_vars[key].set(file_path)
    
    def load_saved_config(self):
        """Load saved configuration into GUI"""
        config = self.config_manager.get_config_dict()
        
        self.username_var.set(config['credentials']['username'])
        self.password_var.set(config['credentials']['password'])
        self.url_var.set(config['credentials']['url'])
        
        for key in self.path_vars:
            self.path_vars[key].set(config['paths'][key])
        
        self.gemini_var.set(config['APIs']['gemini'])
        self.fireworks_var.set(config['APIs']['fireworks'])
        
        for key in self.timeout_vars:
            self.timeout_vars[key].set(str(config['timeouts'][key]))
    
    def save_configuration(self):
        """Save current configuration"""
        # Update config manager
        self.config_manager.config['credentials']['username'] = self.username_var.get()
        self.config_manager.config['credentials']['password'] = self.password_var.get()
        self.config_manager.config['credentials']['url'] = self.url_var.get()
        
        for key in self.path_vars:
            self.config_manager.config['paths'][key] = self.path_vars[key].get()
        
        self.config_manager.config['apis']['gemini'] = self.gemini_var.get()
        self.config_manager.config['apis']['fireworks'] = self.fireworks_var.get()
        
        for key in self.timeout_vars:
            self.config_manager.config['timeouts'][key] = self.timeout_vars[key].get()
        
        # Save to file
        self.config_manager.save_config()
        messagebox.showinfo("Success", "Configuration saved successfully!")
    
    def log_message(self, message):
        """Add message to log"""
        self.log_text.insert(tk.END, f"{message}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()
    
    def validate_configuration(self):
        """Validate that all required fields are filled"""
        config = self.config_manager.get_config_dict()
        
        required_fields = [
            config['credentials']['username'],
            config['credentials']['password'],
            config['paths']['word'],
            config['paths']['excel'],
            config['paths']['done'],
            config['paths']['failed']
        ]
        
        if not all(required_fields):
            return False, "Please fill in all required fields"
        
        # Check if paths exist
        if not os.path.exists(config['paths']['word']):
            return False, f"Word input path does not exist: {config['paths']['word']}"
        
        return True, "Configuration is valid"
    
    def process_articles(self):
        """Main processing function that runs in a separate thread"""
        try:
            self.log_message("🚀 Starting article processing...")
            
            # Validate configuration
            is_valid, message = self.validate_configuration()
            if not is_valid:
                self.log_message(f"❌ Configuration error: {message}")
                return
            
            # Get current configuration
            config = self.config_manager.get_config_dict()
            global CONFIG
            CONFIG = config
            
            # Create necessary directories
            for path_key in ['excel', 'done', 'failed', 'featured_images_path']:
                os.makedirs(config['paths'][path_key], exist_ok=True)
            
            self.log_message("📄 Starting DOCX to Excel conversion...")
            
            # Convert DOCX to Excel
            converter = DocxToExcelConverter()
            converter.convert(config["paths"]["word"], config["paths"]["excel"])
            excel_files = converter.excel_paths
            
            if not excel_files:
                self.log_message("❌ No Excel files were generated from DOCX conversion")
                return
            
            self.log_message(f"✅ Generated {len(excel_files)} Excel files")
            
            # Process each Excel file
            total_files = len(excel_files)
            for i, file in enumerate(excel_files):
                if self.stop_flag.is_set():
                    self.log_message("🛑 Processing stopped by user")
                    break
                
                try:
                    self.log_message(f"📝 Processing article {i+1}/{total_files}: {os.path.basename(file)}")
                    self.status_var.set(f"Processing {i+1}/{total_files}")
                    
                    article_writter(file, config)
                    shutil.move(file, config['paths']['done'])
                    self.log_message(f"✅ Successfully processed: {os.path.basename(file)}")
                    
                except Exception as e:
                    self.log_message(f"❌ Failed to process {os.path.basename(file)}: {str(e)}")
                    try:
                        shutil.move(file, config['paths']['failed'])
                    except:
                        pass
                
                # Update progress
                progress = ((i + 1) / total_files) * 100
                self.progress_var.set(progress)
            
            self.log_message("🎉 Processing completed!")
            self.status_var.set("Completed")
            
        except Exception as e:
            self.log_message(f"❌ Critical error: {str(e)}")
            self.status_var.set("Error occurred")
        
        finally:
            self.progress_var.set(0)
    
    def start_processing(self):
        """Start processing in a separate thread"""
        if self.processing_thread and self.processing_thread.is_alive():
            messagebox.showwarning("Warning", "Processing is already running!")
            return
        
        # Save current configuration before processing
        self.save_configuration()
        
        # Clear log and reset flags
        self.log_text.delete(1.0, tk.END)
        self.stop_flag.clear()
        self.progress_var.set(0)
        
        # Start processing thread
        self.processing_thread = threading.Thread(target=self.process_articles)
        self.processing_thread.daemon = True
        self.processing_thread.start()
    
    def stop_processing(self):
        """Stop the processing"""
        self.stop_flag.set()
        self.status_var.set("Stopping...")
        self.log_message("🛑 Stop signal sent...")

def main():
    root = tk.Tk()
    app = ArticleAutomationGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()