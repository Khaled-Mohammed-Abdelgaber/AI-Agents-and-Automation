import os
import re
import json
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading
import queue
import pandas as pd
from docx import Document
import google.generativeai as genai
from google.generativeai import GenerativeModel
from dotenv import load_dotenv
import undetected_chromedriver as uc
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
import csv
import requests
import random
import shutil
from PIL import Image, ImageEnhance

load_dotenv()

class ConfigManager:
    def __init__(self, config_file="config.json"):
        self.config_file = config_file
        self.default_config = {
            "credentials": {
                "username": "",
                "password": "",
                "url": "https://pharmastan.net/pharma-login"
            },
            "paths": {
                "excel": "",
                "word": "",
                "failed": "",
                "done": "",
                "logo_image_path": "",
                "original_images_path": "",
                "featured_images_path": "",
            },
            "timeouts": {
                "page_load": 60000,
                "element": 10000,
                "hover": 1000,
                "click": 2000,
                "typing": 120
            },
            "APIs": {
                "gemini": "",
                "firworks": ""
            },
            "browser": {
                "headless": False
            }
        }
    
    def load_config(self):
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    loaded_config = json.load(f)
                
                # Validate and merge with default config to ensure all keys exist
                config = self._merge_with_defaults(loaded_config)
                return config
            return self.default_config.copy()
        except (json.JSONDecodeError, FileNotFoundError, PermissionError) as e:
            print(f"Error loading config file: {e}")
            # Try to backup corrupted config
            self._backup_corrupted_config()
            return self.default_config.copy()
        except Exception as e:
            print(f"Unexpected error loading config: {e}")
            return self.default_config.copy()
    
    def _merge_with_defaults(self, loaded_config):
        """Merge loaded config with defaults to ensure all required keys exist"""
        config = self.default_config.copy()
        
        # Recursively merge configurations
        def deep_merge(default, loaded):
            for key, value in loaded.items():
                if key in default:
                    if isinstance(default[key], dict) and isinstance(value, dict):
                        deep_merge(default[key], value)
                    else:
                        default[key] = value
        
        if isinstance(loaded_config, dict):
            deep_merge(config, loaded_config)
        
        return config
    
    def _backup_corrupted_config(self):
        """Create a backup of corrupted config file"""
        try:
            if os.path.exists(self.config_file):
                import time
                backup_name = f"{self.config_file}.backup_{int(time.time())}"
                import shutil
                shutil.copy2(self.config_file, backup_name)
                print(f"Corrupted config backed up to: {backup_name}")
        except Exception as e:
            print(f"Could not backup corrupted config: {e}")
    
    def save_config(self, config):
        try:
            # Validate config structure before saving
            if not self._validate_config_structure(config):
                print("Invalid config structure, using defaults")
                config = self.default_config.copy()
            
            # Create backup of existing config before overwriting
            self._backup_existing_config()
            
            # Save with atomic write (write to temp file first, then rename)
            temp_file = f"{self.config_file}.tmp"
            with open(temp_file, 'w', encoding='utf-8') as f:
                json.dump(config, f, indent=4, ensure_ascii=False)
            
            # Atomic rename (safer than direct write)
            import os
            if os.path.exists(self.config_file):
                os.remove(self.config_file)
            os.rename(temp_file, self.config_file)
            
            return True
        except (PermissionError, OSError) as e:
            print(f"File system error saving config: {e}")
            return False
        except Exception as e:
            print(f"Error saving config: {e}")
            # Clean up temp file if it exists
            try:
                temp_file = f"{self.config_file}.tmp"
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except:
                pass
            return False
    
    def _validate_config_structure(self, config):
        """Validate that config has the required structure"""
        try:
            required_keys = ["credentials", "paths", "timeouts", "APIs", "browser"]
            if not isinstance(config, dict):
                return False
            
            for key in required_keys:
                if key not in config:
                    return False
                if not isinstance(config[key], dict):
                    return False
            
            return True
        except:
            return False
    
    def _backup_existing_config(self):
        """Create a backup of existing config before overwriting"""
        try:
            if os.path.exists(self.config_file):
                backup_name = f"{self.config_file}.bak"
                import shutil
                shutil.copy2(self.config_file, backup_name)
        except Exception as e:
            print(f"Could not backup existing config: {e}")

class DocxToExcelConverter:
    def __init__(self):
        self.paragraphs = []
        self.articles_data = {}
        self.excel_paths = []

    def read_docx(self, docx_path):
        doc = Document(docx_path)
        self.paragraphs = doc.paragraphs

    def get_actual_bullet_or_number(self, paragraph):
        element = paragraph._element
        numPr = element.find('.//w:numPr', namespaces=element.nsmap)
        if numPr is None:
            return None

        bullet_char = '•'
        lvl = element.find('.//w:ilvl', namespaces=element.nsmap)
        if lvl is not None:
            return bullet_char + ' '

        numId = element.find('.//w:numId', namespaces=element.nsmap)
        if numId is not None:
            return "1. "

        return None

    def get_clean_text_with_formatting(self):
        lines = []

        for para in self.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue

            prefix = self.get_actual_bullet_or_number(para)
            if prefix:
                lines.append(prefix + text)
            else:
                pPr = para._element.find('.//w:pPr', namespaces=para._element.nsmap)
                if pPr is not None:
                    ind = pPr.find('.//w:ind', namespaces=para._element.nsmap)
                    if ind is not None:
                        left = ind.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left')
                        if left:
                            lines.append('    ' * (int(left) // 720) + text)
                            continue
                lines.append(text)
        return "\n".join(lines)

    def sanitize_filename(self, name):
        return re.sub(r'[\\/*?:"<>|]', "_", name)

    def extract_structure(self):
        self.articles_data = {}
        full_text = self.get_clean_text_with_formatting()

        article_pattern = r'###\s*(.*?)\s*###'
        article_matches = list(re.finditer(article_pattern, full_text, re.DOTALL))

        for i, article_match in enumerate(article_matches):
            article_title = article_match.group(1).strip()
            start = article_match.end()
            end = article_matches[i + 1].start() if i + 1 < len(article_matches) else len(full_text)
            article_content = full_text[start:end].strip()

            subsection_pattern = r'\${3}\s*(.*?)\s*\${3}'
            subsection_matches = list(re.finditer(subsection_pattern, article_content, re.DOTALL))

            sections = []
            for j, sub_match in enumerate(subsection_matches):
                subsection_title = sub_match.group(1).strip()
                body_start = sub_match.end()
                body_end = subsection_matches[j + 1].start() if j + 1 < len(subsection_matches) else len(article_content)
                subsection_body = article_content[body_start:body_end].strip()
                subsection_body = re.sub(r'\n{2,}', '\n', subsection_body)
                sections.append([article_title, subsection_title, subsection_body])

            self.articles_data[article_title] = sections

    def save_each_article_to_excel(self, output_folder):
        os.makedirs(output_folder, exist_ok=True)
        for article_title, sections in self.articles_data.items():
            safe_title = self.sanitize_filename(article_title)
            output_path = os.path.join(output_folder, f"{safe_title}.xlsx")
            df = pd.DataFrame(sections, columns=["Article Title", "Subsection Title", "Content"])
            self.excel_paths.append(output_path)            
            writer = pd.ExcelWriter(output_path, engine='xlsxwriter')
            df.to_excel(writer, index=False, sheet_name='Content')
            workbook = writer.book
            worksheet = writer.sheets['Content']
            worksheet.set_column('A:A', 30)
            worksheet.set_column('B:B', 30)
            worksheet.set_column('C:C', 80)
            text_format = workbook.add_format({'text_wrap': True})
            worksheet.set_column('C:C', None, text_format)
            writer.close()

    def convert(self, input_folder, output_folder, progress_callback=None):
        if not os.path.exists(input_folder):
            raise FileNotFoundError(f"Input folder not found: {input_folder}")
        
        os.makedirs(output_folder, exist_ok=True)
        files = [f for f in os.listdir(input_folder) if f.lower().endswith('.docx')]
        
        for i, filename in enumerate(files):
            if progress_callback:
                progress_callback(f"Processing: {filename}")
            
            full_path = os.path.join(input_folder, filename)
            self.read_docx(full_path)
            self.extract_structure()
            base_name = os.path.splitext(filename)[0]
            safe_folder = self.sanitize_filename(base_name)
            article_output_folder = os.path.join(output_folder, safe_folder)
            self.save_each_article_to_excel(article_output_folder)
        
        if progress_callback:
            progress_callback("All Word files processed successfully!")

class AutomationEngine:
    def __init__(self, config):
        self.config = config
    
    def convert_to_paragraphs(self, text: str) -> str:
        lines = text.splitlines()
        paragraphs = []
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if re.match(r"^[\-\•\*]\s+", stripped):
                bullet_content = re.sub(r"^[\-\•\*]\s*", "", stripped)
                paragraphs.append(
                    '<p dir="rtl" style="text-indent:-1.5em; padding-right:2em;">'
                    '&bull;&emsp;' + bullet_content +
                    '</p>'
                )
            else:
                paragraphs.append(f'<p dir="rtl">{stripped}</p>')
        return ''.join(paragraphs)

    def generate_arabic_meta_expert_summary_fireworks(self, excel_path):
        url = "https://api.fireworks.ai/inference/v1/completions"
        df = pd.read_excel(excel_path)
        df.fillna("", inplace=True)

        article_title = df.iloc[0, 0]
        article_intro = df.iloc[0, 2]

        subsections = []
        for _, row in df.iterrows():
            subtitle = row.iloc[1]
            body = row.iloc[2]
            if subtitle or body:
                subsections.append(f"{subtitle}\n{body}".strip())
        full_content = f"{article_intro.strip()}\n\n" + "\n\n".join(subsections)

        prompt = (
            f"أنشئ وصفًا ميتا احترافيًا باللغة العربية لا يتجاوز 160 حرفًا،"
            f" دون شرح أو خطوات، ويكون موجزًا ومناسبًا لمحركات البحث (SEO)،"
            f" ويحتوي على كلمات مفتاحية طبية مناسبة لموضوع المقال التالي:\n\n"
            f"العنوان: {article_title}\n\n"
            f"المحتوى:\n{full_content}\n\n"
            f"الوصف الميتا:"
        )
        
        payload = {
            "model": "accounts/fireworks/models/llama4-maverick-instruct-basic",
            "max_tokens": 70,
            "top_p": 1,
            "top_k": 40,
            "temperature": 0.1,
            "presence_penalty": 0,
            "frequency_penalty": 0,
            "stop": ["\n", "انتهى", "الخطوة", "هل يمكن"],
            "prompt": prompt
        }
        
        headers = {
            "Accept": "application/json",
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.config['APIs']['firworks']}"
        }
        
        response = requests.post(url, headers=headers, data=json.dumps(payload))
        if response.status_code == 200:
            result = response.json()
            meta_text = result.get("choices", [{}])[0].get("text", "").strip()
            if not meta_text:
                raise ValueError("Fireworks response is empty")
            return {'meta_expert_summary': meta_text}
        else:
            raise Exception(f"API request failed with status code {response.status_code}")

    def generate_arabic_meta_expert_summary(self, excel_path):
        genai.configure(api_key=self.config['APIs']['gemini'])
        model = GenerativeModel('gemini-2.0-flash-lite-001')
        df = pd.read_excel(excel_path)
        df.fillna("", inplace=True)
        
        article_title = df.iloc[0,0]
        article_intro = df.iloc[0,2]
        
        subsections = []
        for _, row in df.iterrows():
            title = row.iloc[1] 
            body = row.iloc[2] 
            if title or body:
                subsections.append(f"{title}\n{body}")

        full_content = f"{article_intro}\n\n" + "\n\n".join(subsections)

        prompt = f"""
المقال التالي يتحدث عن موضوع طبي بعنوان "{article_title}".

المحتوى:
{full_content}

🧠 المطلوب:
اكتب ملخصًا احترافيًا موجزًا (سطرين فقط كحد أقصى) يصلح كوصف ميتا للمقال. يجب أن يكون باللغة العربية ويعبر بدقة عن محتوى المقال بطريقة جذابة وواضحة.
        """
        
        response = model.generate_content(prompt)
        if response.text.strip() == "":
            raise ValueError("gemini response is empty")

        return {'meta_expert_summary': response.text.strip()}

    def add_watermark_from_folder(self, folder_path, output_path, watermark_image_path, 
                                base_size=(616, 367), transparency=0.3, 
                                allowed_extensions=(".jpg", ".jpeg", ".png")):
        image_files = [
            f for f in os.listdir(folder_path)
            if f.lower().endswith(allowed_extensions)
        ]
        if not image_files:
            raise FileNotFoundError("No image files found in the folder.")
        
        chosen_image = random.choice(image_files)
        base_image_path = os.path.join(folder_path, chosen_image)
        
        base_img = Image.open(base_image_path).convert("RGBA")
        base_img = base_img.resize(base_size)
        
        watermark = Image.open(watermark_image_path).convert("RGBA")
        watermark = watermark.resize(base_size)
        
        alpha = watermark.split()[3]
        alpha = ImageEnhance.Brightness(alpha).enhance(transparency)
        watermark.putalpha(alpha)

        combined = Image.alpha_composite(base_img, watermark)
        
        os.makedirs(output_path, exist_ok=True)
        output_file = os.path.join(output_path, chosen_image)
        combined.convert("RGB").save(output_file, "JPEG")
        return output_file

    # Import automation functions from external file
    from automation_functions import article_writer

class PharmaAutomationApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Pharma Content Automation")
        self.root.geometry("900x700")
        
        # Initialize configuration manager with error handling
        try:
            self.config_manager = ConfigManager()
            self.config = self.config_manager.load_config()
            config_status = "Configuration loaded successfully"
        except Exception as e:
            print(f"Critical error initializing configuration: {e}")
            # Fallback to default configuration
            self.config_manager = ConfigManager()
            self.config = self.config_manager.default_config.copy()
            config_status = "Using default configuration due to load error"
        
        # Create UI
        try:
            self.create_widgets()
            self.load_saved_config()
        except Exception as e:
            print(f"Error creating UI: {e}")
            messagebox.showerror("UI Error", f"Error creating user interface: {e}")
            return
        
        # Queue for thread communication
        self.queue = queue.Queue()
        self.root.after(100, self.process_queue)
        
        # Processing state
        self.excel_files = []
        self.is_processing = False
        
        # Show config status in log after UI is ready
        self.root.after(500, lambda: self.log_message(config_status))
        
        # Auto-detect Excel files on startup
        self.root.after(1000, self.auto_detect_excel_on_startup)

    def create_widgets(self):
        # Create notebook for tabs
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Configuration tab
        self.config_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.config_frame, text="Configuration")
        self.create_config_tab()
        
        # Process tab
        self.process_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.process_frame, text="Process Articles")
        self.create_process_tab()

    def create_config_tab(self):
        # Main container with scrollbar
        canvas = tk.Canvas(self.config_frame)
        scrollbar = ttk.Scrollbar(self.config_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Credentials section
        cred_frame = ttk.LabelFrame(scrollable_frame, text="Website Credentials", padding=10)
        cred_frame.pack(fill="x", padx=10, pady=5)
        
        ttk.Label(cred_frame, text="Username:").grid(row=0, column=0, sticky="w", padx=5, pady=2)
        self.username_var = tk.StringVar()
        ttk.Entry(cred_frame, textvariable=self.username_var, width=50).grid(row=0, column=1, padx=5, pady=2)
        
        ttk.Label(cred_frame, text="Password:").grid(row=1, column=0, sticky="w", padx=5, pady=2)
        self.password_var = tk.StringVar()
        ttk.Entry(cred_frame, textvariable=self.password_var, show="*", width=50).grid(row=1, column=1, padx=5, pady=2)
        
        ttk.Label(cred_frame, text="Login URL:").grid(row=2, column=0, sticky="w", padx=5, pady=2)
        self.url_var = tk.StringVar()
        ttk.Entry(cred_frame, textvariable=self.url_var, width=50).grid(row=2, column=1, padx=5, pady=2)
        
        # Paths section
        paths_frame = ttk.LabelFrame(scrollable_frame, text="File Paths", padding=10)
        paths_frame.pack(fill="x", padx=10, pady=5)
        
        self.path_vars = {}
        path_labels = [
            ("Word Documents Folder:", "word"),
            ("Excel Output Folder:", "excel"),
            ("Failed Files Folder:", "failed"),
            ("Completed Files Folder:", "done"),
            ("Logo Image Path:", "logo_image_path"),
            ("Original Images Folder:", "original_images_path"),
            ("Featured Images Folder:", "featured_images_path")
        ]
        
        for i, (label, key) in enumerate(path_labels):
            ttk.Label(paths_frame, text=label).grid(row=i, column=0, sticky="w", padx=5, pady=2)
            self.path_vars[key] = tk.StringVar()
            entry_frame = ttk.Frame(paths_frame)
            entry_frame.grid(row=i, column=1, padx=5, pady=2, sticky="ew")
            
            ttk.Entry(entry_frame, textvariable=self.path_vars[key], width=40).pack(side="left", fill="x", expand=True)
            if "image_path" in key and key != "logo_image_path":
                ttk.Button(entry_frame, text="Browse", 
                          command=lambda k=key: self.browse_folder(k)).pack(side="right", padx=(5,0))
            elif key == "logo_image_path":
                ttk.Button(entry_frame, text="Browse", 
                          command=lambda k=key: self.browse_file(k)).pack(side="right", padx=(5,0))
            else:
                ttk.Button(entry_frame, text="Browse", 
                          command=lambda k=key: self.browse_folder(k)).pack(side="right", padx=(5,0))
        
        paths_frame.columnconfigure(1, weight=1)
        
        # API Keys section
        api_frame = ttk.LabelFrame(scrollable_frame, text="API Keys", padding=10)
        api_frame.pack(fill="x", padx=10, pady=5)
        
        ttk.Label(api_frame, text="Gemini API Key:").grid(row=0, column=0, sticky="w", padx=5, pady=2)
        self.gemini_api_var = tk.StringVar()
        ttk.Entry(api_frame, textvariable=self.gemini_api_var, width=50, show="*").grid(row=0, column=1, padx=5, pady=2)
        
        ttk.Label(api_frame, text="Fireworks API Key:").grid(row=1, column=0, sticky="w", padx=5, pady=2)
        self.firworks_api_var = tk.StringVar()
        ttk.Entry(api_frame, textvariable=self.firworks_api_var, width=50, show="*").grid(row=1, column=1, padx=5, pady=2)
        
        # Browser Settings section
        browser_frame = ttk.LabelFrame(scrollable_frame, text="Browser Settings", padding=10)
        browser_frame.pack(fill="x", padx=10, pady=5)
        
        # Headless mode checkbox
        self.headless_var = tk.BooleanVar()
        headless_checkbox = ttk.Checkbutton(
            browser_frame, 
            text="Run browser in headless mode (hidden browser window)", 
            variable=self.headless_var,
            command=self.on_headless_change
        )
        headless_checkbox.pack(anchor="w", padx=5, pady=5)
        
        # Info label
        info_label = ttk.Label(
            browser_frame, 
            text="💡 Headless mode hides the browser window during automation.\nDisable for debugging or to watch the process.",
            foreground="gray"
        )
        info_label.pack(anchor="w", padx=5, pady=(0, 5))
        
        # Buttons
        button_frame = ttk.Frame(scrollable_frame)
        button_frame.pack(fill="x", padx=10, pady=10)
        
        ttk.Button(button_frame, text="Save Configuration", command=self.save_config).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Load Configuration", command=self.load_config).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Reset to Default", command=self.reset_config).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Fix Config Issues", command=self.fix_config_issues).pack(side="left", padx=5)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

    def create_process_tab(self):
        # Status section
        status_frame = ttk.LabelFrame(self.process_frame, text="Process Status", padding=10)
        status_frame.pack(fill="x", padx=10, pady=5)
        
        # File status
        file_status_frame = ttk.Frame(status_frame)
        file_status_frame.pack(fill="x", pady=(0, 10))
        
        ttk.Label(file_status_frame, text="Word Documents:").pack(side="left")
        self.word_count_var = tk.StringVar(value="0 files")
        ttk.Label(file_status_frame, textvariable=self.word_count_var, foreground="blue").pack(side="left", padx=(5, 20))
        
        ttk.Label(file_status_frame, text="Excel Files Ready:").pack(side="left")
        self.excel_count_var = tk.StringVar(value="0 files")
        ttk.Label(file_status_frame, textvariable=self.excel_count_var, foreground="green").pack(side="left", padx=(5, 20))
        
        # Browser mode indicator
        ttk.Label(file_status_frame, text="Browser Mode:").pack(side="left")
        self.browser_mode_var = tk.StringVar(value="Visible")
        ttk.Label(file_status_frame, textvariable=self.browser_mode_var, foreground="purple").pack(side="left", padx=(5, 20))
        
        # Buttons for file management
        buttons_frame = ttk.Frame(file_status_frame)
        buttons_frame.pack(side="right")
        
        ttk.Button(buttons_frame, text="Show Details", command=self.show_excel_details).pack(side="right", padx=(5,0))
        ttk.Button(buttons_frame, text="Scan Excel", command=self.scan_excel_files).pack(side="right", padx=(5,0))
        ttk.Button(buttons_frame, text="Clear List", command=self.clear_excel_files_list).pack(side="right", padx=(5,0))
        ttk.Button(buttons_frame, text="Refresh", command=self.refresh_file_counts).pack(side="right", padx=(5,0))
        
        # Progress bar
        self.progress_var = tk.StringVar(value="Ready to start processing...")
        ttk.Label(status_frame, textvariable=self.progress_var).pack(anchor="w")
        
        self.progress_bar = ttk.Progressbar(status_frame, mode='indeterminate')
        self.progress_bar.pack(fill="x", pady=5)
        
        # Log area
        log_frame = ttk.LabelFrame(self.process_frame, text="Processing Log", padding=10)
        log_frame.pack(fill="both", expand=True, padx=10, pady=5)
        
        self.log_text = scrolledtext.ScrolledText(log_frame, height=15, state='disabled')
        self.log_text.pack(fill="both", expand=True)
        
        # Control buttons
        control_frame = ttk.Frame(self.process_frame)
        control_frame.pack(fill="x", padx=10, pady=5)
        
        # Left side buttons (main actions)
        left_buttons = ttk.Frame(control_frame)
        left_buttons.pack(side="left", fill="x", expand=True)
        
        self.convert_button = ttk.Button(left_buttons, text="Convert Word to Excel", 
                                       command=self.start_conversion, style="Accent.TButton")
        self.convert_button.pack(side="left", padx=5)
        
        self.publish_button = ttk.Button(left_buttons, text="Publish to WordPress", 
                                       command=self.start_publication, state="disabled")
        self.publish_button.pack(side="left", padx=5)
        
        # Right side buttons (utility)
        right_buttons = ttk.Frame(control_frame)
        right_buttons.pack(side="right")
        
        self.stop_button = ttk.Button(right_buttons, text="Stop", 
                                    command=self.stop_processing, state="disabled")
        self.stop_button.pack(side="right", padx=5)
        
        ttk.Button(right_buttons, text="Clear Log", command=self.clear_log).pack(side="right", padx=5)

    def browse_folder(self, key):
        folder = filedialog.askdirectory()
        if folder:
            self.path_vars[key].set(folder)

    def browse_file(self, key):
        file = filedialog.askopenfilename(
            filetypes=[("Image files", "*.png *.jpg *.jpeg *.gif *.bmp")]
        )
        if file:
            self.path_vars[key].set(file)

    def on_headless_change(self):
        """Update browser mode indicator when headless checkbox changes"""
        self.browser_mode_var.set("Headless" if self.headless_var.get() else "Visible")

    def save_config(self):
        try:
            # Update config with current values
            # Ensure config structure exists
            if "credentials" not in self.config:
                self.config["credentials"] = {}
            if "paths" not in self.config:
                self.config["paths"] = {}
            if "APIs" not in self.config:
                self.config["APIs"] = {}
            if "browser" not in self.config:
                self.config["browser"] = {}
            
            # Update values safely
            self.config["credentials"]["username"] = self.username_var.get()
            self.config["credentials"]["password"] = self.password_var.get()
            self.config["credentials"]["url"] = self.url_var.get()
            
            for key, var in self.path_vars.items():
                self.config["paths"][key] = var.get()
            
            self.config["APIs"]["gemini"] = self.gemini_api_var.get()
            self.config["APIs"]["firworks"] = self.firworks_api_var.get()
            
            self.config["browser"]["headless"] = self.headless_var.get()
            
            # Attempt to save
            if self.config_manager.save_config(self.config):
                messagebox.showinfo("Success", "Configuration saved successfully!")
                self.log_message("Configuration saved successfully")
            else:
                messagebox.showerror("Error", "Failed to save configuration! Check file permissions.")
                self.log_message("❌ Failed to save configuration")
                
        except Exception as e:
            error_msg = f"Error preparing configuration for save: {e}"
            print(error_msg)
            messagebox.showerror("Error", f"Failed to save configuration: {str(e)}")
            self.log_message(f"❌ {error_msg}")

    def load_config(self):
        self.config = self.config_manager.load_config()
        self.load_saved_config()
        messagebox.showinfo("Success", "Configuration loaded successfully!")

    def load_saved_config(self):
        # Load values into UI with safe fallbacks
        try:
            # Load credentials
            credentials = self.config.get("credentials", {})
            self.username_var.set(credentials.get("username", ""))
            self.password_var.set(credentials.get("password", ""))
            self.url_var.set(credentials.get("url", "https://pharmastan.net/pharma-login"))
            
            # Load paths
            paths = self.config.get("paths", {})
            for key, var in self.path_vars.items():
                var.set(paths.get(key, ""))
            
            # Load API keys
            apis = self.config.get("APIs", {})
            self.gemini_api_var.set(apis.get("gemini", ""))
            self.firworks_api_var.set(apis.get("firworks", ""))
            
            # Load browser settings
            browser = self.config.get("browser", {})
            self.headless_var.set(browser.get("headless", False))
            
        except Exception as e:
            print(f"Error loading config values into UI: {e}")
            # Set all values to defaults if there's an error
            self.username_var.set("")
            self.password_var.set("")
            self.url_var.set("https://pharmastan.net/pharma-login")
            
            for key, var in self.path_vars.items():
                var.set("")
            
            self.gemini_api_var.set("")
            self.firworks_api_var.set("")
            self.headless_var.set(False)
            
            messagebox.showwarning("Configuration Warning", 
                                 "Error loading some configuration values. Please check and save your settings.")
        
        # Refresh file counts after loading config
        self.root.after(100, self.refresh_file_counts)

    def reset_config(self):
        if messagebox.askyesno("Confirm Reset", "Are you sure you want to reset to default configuration?"):
            self.config = self.config_manager.default_config.copy()
            self.load_saved_config()
            self.log_message("Configuration reset to defaults")

    def fix_config_issues(self):
        """Attempt to fix common configuration issues"""
        try:
            # Check if config file exists and is readable
            config_file = self.config_manager.config_file
            issues_found = []
            fixes_applied = []
            
            if os.path.exists(config_file):
                try:
                    with open(config_file, 'r', encoding='utf-8') as f:
                        test_config = json.load(f)
                except json.JSONDecodeError:
                    issues_found.append("Corrupted JSON in config file")
                    # Create backup and reset
                    self.config_manager._backup_corrupted_config()
                    self.config = self.config_manager.default_config.copy()
                    fixes_applied.append("Created backup of corrupted config and reset to defaults")
                except Exception as e:
                    issues_found.append(f"Cannot read config file: {e}")
            
            # Validate current config structure
            if not self.config_manager._validate_config_structure(self.config):
                issues_found.append("Invalid configuration structure")
                self.config = self.config_manager.default_config.copy()
                fixes_applied.append("Reset configuration structure to defaults")
            
            # Check file permissions
            try:
                test_file = f"{config_file}.test"
                with open(test_file, 'w') as f:
                    f.write("test")
                os.remove(test_file)
            except Exception as e:
                issues_found.append(f"File permission issues: {e}")
                fixes_applied.append("Please check file permissions manually")
            
            # Update UI with current config
            self.load_saved_config()
            
            # Show results
            if issues_found:
                message = "Issues found:\n" + "\n".join(f"• {issue}" for issue in issues_found)
                if fixes_applied:
                    message += "\n\nFixes applied:\n" + "\n".join(f"• {fix}" for fix in fixes_applied)
                messagebox.showinfo("Configuration Check", message)
                self.log_message(f"Fixed {len(fixes_applied)} configuration issues")
            else:
                messagebox.showinfo("Configuration Check", "No configuration issues found!")
                self.log_message("Configuration check passed - no issues found")
                
        except Exception as e:
            error_msg = f"Error checking configuration: {e}"
            messagebox.showerror("Error", error_msg)
            self.log_message(f"❌ {error_msg}")

    def log_message(self, message):
        self.log_text.config(state='normal')
        self.log_text.insert(tk.END, f"{time.strftime('%H:%M:%S')} - {message}\n")
        self.log_text.config(state='disabled')
        self.log_text.see(tk.END)

    def clear_log(self):
        self.log_text.config(state='normal')
        self.log_text.delete(1.0, tk.END)
        self.log_text.config(state='disabled')

    def refresh_file_counts(self):
        # Count Word documents
        word_count = 0
        if self.path_vars["word"].get() and os.path.exists(self.path_vars["word"].get()):
            word_files = [f for f in os.listdir(self.path_vars["word"].get()) if f.lower().endswith('.docx')]
            word_count = len(word_files)
        
        self.word_count_var.set(f"{word_count} files")
        
        # Detect existing Excel files in the Excel output folder
        self.detect_existing_excel_files()
        
        # Count Excel files ready for publication
        excel_count = len(self.excel_files)
        self.excel_count_var.set(f"{excel_count} files")
        
        # Update button states
        if excel_count > 0:
            self.publish_button.config(state="normal" if not self.is_processing else "disabled")
        else:
            self.publish_button.config(state="disabled")
        
        # Update browser mode indicator
        headless_mode = self.config.get("browser", {}).get("headless", False)
        self.browser_mode_var.set("Headless" if headless_mode else "Visible")

    def detect_existing_excel_files(self):
        """Detect existing Excel files in the Excel output folder"""
        try:
            excel_folder = self.path_vars["excel"].get()
            if not excel_folder or not os.path.exists(excel_folder):
                return
            
            detected_files = []
            done_folder = self.path_vars["done"].get()
            failed_folder = self.path_vars["failed"].get()
            
            # Walk through all subdirectories in the Excel folder
            for root, dirs, files in os.walk(excel_folder):
                for file in files:
                    if file.lower().endswith('.xlsx'):
                        full_path = os.path.join(root, file)
                        
                        # Only add if it's not already in our list
                        if full_path not in self.excel_files:
                            # Check if this file was already processed (exists in done/failed folders)
                            is_already_processed = False
                            
                            if done_folder and os.path.exists(done_folder):
                                done_file_path = os.path.join(done_folder, file)
                                if os.path.exists(done_file_path):
                                    is_already_processed = True
                            
                            if failed_folder and os.path.exists(failed_folder):
                                failed_file_path = os.path.join(failed_folder, file)
                                if os.path.exists(failed_file_path):
                                    is_already_processed = True
                            
                            if not is_already_processed:
                                detected_files.append(full_path)
                            else:
                                self.log_message(f"⏩ Skipping already processed file: {file}")
            
            # Add detected files to the excel_files list
            if detected_files:
                self.excel_files.extend(detected_files)
                self.log_message(f"🔍 Detected {len(detected_files)} existing Excel files ready for publication")
                
                # Log the detected files
                for file_path in detected_files:
                    relative_path = os.path.relpath(file_path, excel_folder)
                    self.log_message(f"  → Found: {relative_path}")
                    
        except Exception as e:
            self.log_message(f"❌ Error detecting Excel files: {str(e)}")
            print(f"Error detecting Excel files: {e}")

    def scan_excel_files(self):
        """Manually scan for Excel files and show results"""
        try:
            excel_folder = self.path_vars["excel"].get()
            if not excel_folder:
                messagebox.showwarning("No Excel Folder", "Please set the Excel Output Folder path first.")
                return
            
            if not os.path.exists(excel_folder):
                messagebox.showwarning("Folder Not Found", f"Excel folder does not exist:\n{excel_folder}")
                return
            
            # Clear current Excel files list
            old_count = len(self.excel_files)
            self.excel_files.clear()
            
            # Detect all Excel files
            self.detect_existing_excel_files()
            
            new_count = len(self.excel_files)
            
            # Update UI
            self.refresh_file_counts()
            
            # Show results
            if new_count > 0:
                messagebox.showinfo("Excel Files Found", 
                                  f"Found {new_count} Excel files ready for publication!\n\n"
                                  f"You can now use 'Publish to WordPress' without converting Word documents.")
                self.log_message(f"✅ Manual scan completed: {new_count} Excel files found")
            else:
                messagebox.showinfo("No Excel Files", 
                                  f"No Excel files found in:\n{excel_folder}\n\n"
                                  f"Convert Word documents first or check the Excel folder path.")
                self.log_message("ℹ️ Manual scan completed: No Excel files found")
                
        except Exception as e:
            error_msg = f"Error scanning Excel files: {str(e)}"
            messagebox.showerror("Scan Error", error_msg)
            self.log_message(f"❌ {error_msg}")

    def clear_excel_files_list(self):
        """Clear the current Excel files list"""
        if self.excel_files:
            count = len(self.excel_files)
            self.excel_files.clear()
            self.refresh_file_counts()
            self.log_message(f"🗑️ Cleared {count} Excel files from publication list")
            messagebox.showinfo("List Cleared", f"Cleared {count} Excel files from the publication list.")
        else:
            messagebox.showinfo("List Empty", "No Excel files in the publication list to clear.")

    def auto_detect_excel_on_startup(self):
        """Automatically detect Excel files when the application starts"""
        try:
            excel_folder = self.path_vars["excel"].get()
            if excel_folder and os.path.exists(excel_folder):
                # Check if there are any Excel files
                excel_files_exist = False
                for root, dirs, files in os.walk(excel_folder):
                    for file in files:
                        if file.lower().endswith('.xlsx'):
                            excel_files_exist = True
                            break
                    if excel_files_exist:
                        break
                
                if excel_files_exist:
                    # Clear existing list and detect
                    self.excel_files.clear()
                    self.detect_existing_excel_files()
                    self.refresh_file_counts()
                    
                    if self.excel_files:
                        self.log_message(f"🚀 Startup scan: Found {len(self.excel_files)} Excel files ready for publication")
                else:
                    self.log_message("ℹ️ Startup scan: No Excel files found in output folder")
            else:
                self.log_message("ℹ️ Excel output folder not configured or doesn't exist")
        except Exception as e:
            self.log_message(f"❌ Error in startup Excel scan: {str(e)}")

    def show_excel_details(self):
        """Show details about detected Excel files"""
        if not self.excel_files:
            messagebox.showinfo("No Excel Files", "No Excel files are currently loaded for publication.\n\nUse 'Scan Excel' to detect existing files or 'Convert Word to Excel' to create new ones.")
            return
        
        # Create details window
        details_window = tk.Toplevel(self.root)
        details_window.title("Excel Files Details")
        details_window.geometry("800x500")
        details_window.transient(self.root)
        details_window.grab_set()
        
        # Create scrolled text widget
        text_frame = ttk.Frame(details_window)
        text_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        text_widget = scrolledtext.ScrolledText(text_frame, wrap=tk.WORD, height=20)
        text_widget.pack(fill="both", expand=True)
        
        # Populate with file details
        excel_folder = self.path_vars["excel"].get()
        content = f"Excel Files Ready for Publication ({len(self.excel_files)} files):\n"
        content += "=" * 60 + "\n\n"
        
        for i, file_path in enumerate(self.excel_files, 1):
            try:
                # Get relative path for display
                if excel_folder:
                    relative_path = os.path.relpath(file_path, excel_folder)
                else:
                    relative_path = os.path.basename(file_path)
                
                # Get file info
                file_size = os.path.getsize(file_path)
                file_size_mb = file_size / (1024 * 1024)
                modified_time = os.path.getmtime(file_path)
                modified_date = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(modified_time))
                
                content += f"{i}. {relative_path}\n"
                content += f"   Size: {file_size_mb:.2f} MB\n"
                content += f"   Modified: {modified_date}\n"
                content += f"   Full Path: {file_path}\n\n"
                
            except Exception as e:
                content += f"{i}. {os.path.basename(file_path)} (Error reading details: {e})\n\n"
        
        text_widget.insert(tk.END, content)
        text_widget.config(state='disabled')
        
        # Add close button
        button_frame = ttk.Frame(details_window)
        button_frame.pack(fill="x", padx=10, pady=(0, 10))
        
        ttk.Button(button_frame, text="Close", command=details_window.destroy).pack(side="right")
        ttk.Button(button_frame, text="Refresh List", command=lambda: [self.scan_excel_files(), details_window.destroy()]).pack(side="right", padx=(0, 5))

    def start_conversion(self):
        # Validate basic configuration for conversion
        if not self.validate_conversion_config():
            return
        
        self.convert_button.config(state="disabled")
        self.publish_button.config(state="disabled")
        self.stop_button.config(state="normal")
        self.progress_bar.start()
        self.progress_var.set("Starting Word to Excel conversion...")
        self.is_processing = True
        
        # Start conversion in separate thread
        self.processing_thread = threading.Thread(target=self.convert_documents)
        self.processing_thread.daemon = True
        self.processing_thread.start()

    def start_publication(self):
        # Validate full configuration for publication
        if not self.validate_publication_config():
            return
        
        if not self.excel_files:
            messagebox.showwarning("No Files", "No Excel files available for publication. Please convert Word documents first.")
            return
        
        self.convert_button.config(state="disabled")
        self.publish_button.config(state="disabled")
        self.stop_button.config(state="normal")
        self.progress_bar.start()
        self.progress_var.set("Starting WordPress publication...")
        self.is_processing = True
        
        # Start publication in separate thread
        self.processing_thread = threading.Thread(target=self.publish_articles)
        self.processing_thread.daemon = True
        self.processing_thread.start()

    def stop_processing(self):
        self.convert_button.config(state="normal")
        if self.excel_files:
            self.publish_button.config(state="normal")
        self.stop_button.config(state="disabled")
        self.progress_bar.stop()
        self.progress_var.set("Processing stopped by user")
        self.is_processing = False

    def validate_conversion_config(self):
        # Check required fields for conversion only
        required_fields = [
            (self.path_vars["word"].get(), "Word Documents Folder"),
            (self.path_vars["excel"].get(), "Excel Output Folder"),
        ]
        
        for value, name in required_fields:
            if not value.strip():
                messagebox.showerror("Configuration Error", f"{name} is required for conversion!")
                return False
        
        # Check if word folder exists
        if not os.path.exists(self.path_vars["word"].get()):
            messagebox.showerror("Configuration Error", "Word Documents Folder does not exist!")
            return False
        
        return True

    def validate_publication_config(self):
        # Check required fields for publication
        required_fields = [
            (self.username_var.get(), "Username"),
            (self.password_var.get(), "Password"),
            (self.url_var.get(), "Login URL"),
        ]
        
        for value, name in required_fields:
            if not value.strip():
                messagebox.showerror("Configuration Error", f"{name} is required for publication!")
                return False
        
        # Check API keys
        if not self.gemini_api_var.get().strip() and not self.firworks_api_var.get().strip():
            messagebox.showerror("Configuration Error", "At least one API key (Gemini or Fireworks) is required for publication!")
            return False
        
        # Check image paths if they are provided
        if self.path_vars["logo_image_path"].get() and not os.path.exists(self.path_vars["logo_image_path"].get()):
            messagebox.showerror("Configuration Error", "Logo image path does not exist!")
            return False
        
        if self.path_vars["original_images_path"].get() and not os.path.exists(self.path_vars["original_images_path"].get()):
            messagebox.showerror("Configuration Error", "Original images folder does not exist!")
            return False
        
        return True

    def convert_documents(self):
        try:
            self.queue.put(("log", "Starting Word to Excel conversion..."))
            
            # Save current configuration
            self.save_config()
            
            # Initialize converter
            converter = DocxToExcelConverter()
            
            def progress_callback(message):
                self.queue.put(("log", message))
                self.queue.put(("progress", message))
            
            # Convert Word documents to Excel
            converter.convert(
                self.config["paths"]["word"], 
                self.config["paths"]["excel"],
                progress_callback
            )
            
            # Store the excel files for publication
            self.excel_files = converter.excel_paths
            
            # Move original word files to done folder if conversion is successful
            if self.config["paths"]["done"]:
                os.makedirs(self.config["paths"]["done"], exist_ok=True)
                word_folder = self.config["paths"]["word"]
                for filename in os.listdir(word_folder):
                    if filename.lower().endswith('.docx'):
                        full_path = os.path.join(word_folder, filename)
                        try:
                            shutil.move(full_path, self.config["paths"]["done"])
                            self.queue.put(("log", f"Moved {filename} to done folder"))
                        except Exception as e:
                            self.queue.put(("log", f"Failed to move {filename}: {str(e)}"))
            
            self.queue.put(("log", f"✅ Successfully converted {len(self.excel_files)} documents to Excel!"))
            self.queue.put(("conversion_complete", f"Conversion completed - {len(self.excel_files)} files ready"))
            
        except Exception as e:
            self.queue.put(("log", f"❌ Error during conversion: {str(e)}"))
            self.queue.put(("conversion_complete", "Conversion failed"))

    def publish_articles(self):
        try:
            self.queue.put(("log", "Starting WordPress publication..."))
            
            # Save current configuration
            self.save_config()
            
            # Initialize automation engine
            automation = AutomationEngine(self.config)
            
            published_count = 0
            failed_count = 0
            
            # Process each Excel file
            for excel_file in self.excel_files:
                try:
                    self.queue.put(("log", f"Publishing article: {os.path.basename(excel_file)}"))
                    
                    def article_progress_callback(message):
                        self.queue.put(("log", f"  → {message}"))
                    
                    # Import the article_writer function
                    from automation_functions import article_writer
                    
                    # Call the article writer function
                    article_writer(excel_file, self.config, automation, article_progress_callback)
                    
                    # Move processed file to done folder
                    if self.config["paths"]["done"]:
                        os.makedirs(self.config["paths"]["done"], exist_ok=True)
                        try:
                            shutil.move(excel_file, self.config["paths"]["done"])
                        except:
                            pass  # File might already be moved
                    
                    published_count += 1
                    self.queue.put(("log", f"✅ Successfully published: {os.path.basename(excel_file)}"))
                    
                except Exception as e:
                    failed_count += 1
                    self.queue.put(("log", f"❌ Failed to publish {os.path.basename(excel_file)}: {str(e)}"))
                    
                    # Move failed file to failed folder
                    if self.config["paths"]["failed"]:
                        os.makedirs(self.config["paths"]["failed"], exist_ok=True)
                        try:
                            shutil.move(excel_file, self.config["paths"]["failed"])
                        except:
                            pass
            
            # Clear the excel files list after processing
            self.excel_files = []
            
            self.queue.put(("log", f"✅ Publication completed! Published: {published_count}, Failed: {failed_count}"))
            self.queue.put(("publication_complete", f"Publication completed - {published_count} published, {failed_count} failed"))
            
        except Exception as e:
            self.queue.put(("log", f"❌ Error during publication: {str(e)}"))
            self.queue.put(("publication_complete", "Publication failed"))

    def process_queue(self):
        try:
            while True:
                message_type, message = self.queue.get_nowait()
                
                if message_type == "log":
                    self.log_message(message)
                elif message_type == "progress":
                    self.progress_var.set(message)
                elif message_type == "conversion_complete":
                    self.convert_button.config(state="normal")
                    if self.excel_files:
                        self.publish_button.config(state="normal")
                    self.stop_button.config(state="disabled")
                    self.progress_bar.stop()
                    self.progress_var.set(message)
                    self.is_processing = False
                    self.refresh_file_counts()
                elif message_type == "publication_complete":
                    self.convert_button.config(state="normal")
                    self.publish_button.config(state="disabled")  # No files left to publish
                    self.stop_button.config(state="disabled")
                    self.progress_bar.stop()
                    self.progress_var.set(message)
                    self.is_processing = False
                    self.refresh_file_counts()
                
        except queue.Empty:
            pass
        
        self.root.after(100, self.process_queue)

def main():
    root = tk.Tk()
    
    # Set theme
    style = ttk.Style()
    style.theme_use('clam')
    
    app = PharmaAutomationApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()